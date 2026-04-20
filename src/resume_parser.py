import os
import pdfplumber
from docx import Document
from typing import Optional, Dict, List
import re


class ResumeParser:
    EMAIL_REGEX = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    PHONE_REGEX = r'(\+7|8)[\s\-]?\(?[0-9]{3}\)?[\s\-]?[0-9]{3}[\s\-]?[0-9]{2}[\s\-]?[0-9]{2}'
    
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.raw_text = ""
        self.file_name = os.path.basename(file_path)
    
    def extract_text(self) -> str:
        if self.file_path.endswith('.pdf'):
            self.raw_text = self._parse_pdf()
        elif self.file_path.endswith('.docx'):
            self.raw_text = self._parse_docx()
        elif self.file_path.endswith('.txt'):
            self.raw_text = self._parse_txt()
        else:
            raise ValueError(f"Неподдерживаемый формат: {self.file_path}")
        
        return self.raw_text
    
    def _parse_pdf(self) -> str:
        text_parts = []
        try:
            with pdfplumber.open(self.file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            return '\n'.join(text_parts)
        except Exception as e:
            raise Exception(f"Ошибка парсинга PDF: {str(e)}")
    
    def _parse_docx(self) -> str:
        try:
            doc = Document(self.file_path)
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            return '\n'.join(text_parts)
        except Exception as e:
            raise Exception(f"Ошибка парсинга DOCX: {str(e)}")
    
    def _parse_txt(self) -> str:
        try:
            with open(self.file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            with open(self.file_path, 'r', encoding='cp1251') as f:
                return f.read()
    
    def extract_contacts(self) -> Dict[str, Optional[str]]:
        if not self.raw_text:
            self.extract_text()
        
        contacts = {
            'email': None,
            'phone': None
        }
        
        email_match = re.search(self.EMAIL_REGEX, self.raw_text, re.IGNORECASE)
        if email_match:
            contacts['email'] = email_match.group(0)
        
        phone_match = re.search(self.PHONE_REGEX, self.raw_text)
        if phone_match:
            contacts['phone'] = phone_match.group(0)
        
        return contacts
    
    def extract_name(self) -> Optional[str]:
        if not self.raw_text:
            self.extract_text()
        
        lines = self.raw_text.strip().split('\n')
        
        for line in lines[:5]:
            line = line.strip()
            if line and not re.search(r'[0-9@.]', line) and len(line.split()) <= 3:
                words = line.split()
                if all(w[0].isupper() for w in words if w):
                    return line
        
        return None
    
    def get_summary(self) -> str:
        if not self.raw_text:
            self.extract_text()
        
        return self.raw_text[:2000] if len(self.raw_text) > 2000 else self.raw_text