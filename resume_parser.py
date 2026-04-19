"""
Модуль для извлечения текста из файлов резюме (PDF, DOCX).
"""

import os
import pdfplumber
from docx import Document
from typing import Optional, Dict, List
import re


class ResumeParser:
    """
    Класс для парсинга резюме из PDF и DOCX файлов.
    """
    
    # Регулярные выражения для извлечения контактов
    EMAIL_REGEX = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    PHONE_REGEX = r'(\+7|8)[\s\-]?\(?[0-9]{3}\)?[\s\-]?[0-9]{3}[\s\-]?[0-9]{2}[\s\-]?[0-9]{2}'
    
    def __init__(self, file_path: str):
        """
        Инициализация парсера.
        
        Args:
            file_path: Путь к файлу резюме
        """
        self.file_path = file_path
        self.raw_text = ""
        self.file_name = os.path.basename(file_path)
    
    def extract_text(self) -> str:
        """
        Извлекает текст из файла в зависимости от расширения.
        
        Returns:
            str: Извлеченный текст
        """
        if self.file_path.endswith('.pdf'):
            self.raw_text = self._parse_pdf()
        elif self.file_path.endswith('.docx'):
            self.raw_text = self._parse_docx()
        elif self.file_path.endswith('.txt'):
            self.raw_text = self._parse_txt()
        else:
            raise ValueError(f"Неподдерживаемый формат: {self.file_path}")
        
        return self.raw_text
    
    def _parse_pdf(self) -> str:
        """Парсинг PDF файла."""
        text_parts = []
        try:
            with pdfplumber.open(self.file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            return '\n'.join(text_parts)
        except Exception as e:
            raise Exception(f"Ошибка парсинга PDF: {str(e)}")
    
    def _parse_docx(self) -> str:
        """Парсинг DOCX файла."""
        try:
            doc = Document(self.file_path)
            text_parts = []
            
            # Извлечение текста из параграфов
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Извлечение текста из таблиц
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            return '\n'.join(text_parts)
        except Exception as e:
            raise Exception(f"Ошибка парсинга DOCX: {str(e)}")
    
    def _parse_txt(self) -> str:
        """Парсинг TXT файла."""
        try:
            with open(self.file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            with open(self.file_path, 'r', encoding='cp1251') as f:
                return f.read()
    
    def extract_contacts(self) -> Dict[str, Optional[str]]:
        """
        Извлекает email и телефон из текста.
        
        Returns:
            Dict: Словарь с email и телефоном
        """
        if not self.raw_text:
            self.extract_text()
        
        contacts = {
            'email': None,
            'phone': None
        }
        
        # Поиск email
        email_match = re.search(self.EMAIL_REGEX, self.raw_text, re.IGNORECASE)
        if email_match:
            contacts['email'] = email_match.group(0)
        
        # Поиск телефона
        phone_match = re.search(self.PHONE_REGEX, self.raw_text)
        if phone_match:
            contacts['phone'] = phone_match.group(0)
        
        return contacts
    
    def extract_name(self) -> Optional[str]:
        """
        Пытается извлечь имя кандидата из первых строк резюме.
        """
        if not self.raw_text:
            self.extract_text()
        
        lines = self.raw_text.strip().split('\n')
        
        # Обычно имя находится в первых 5 строках
        for line in lines[:5]:
            line = line.strip()
            # Фильтруем строки без цифр и спецсимволов
            if line and not re.search(r'[0-9@.]', line) and len(line.split()) <= 3:
                # Проверяем, что это похоже на имя (каждое слово с большой буквы)
                words = line.split()
                if all(w[0].isupper() for w in words if w):
                    return line
        
        return None
    
    def get_summary(self) -> str:
        """
        Возвращает краткое содержимое резюме (первые 1000 символов).
        """
        if not self.raw_text:
            self.extract_text()
        
        return self.raw_text[:2000] if len(self.raw_text) > 2000 else self.raw_text